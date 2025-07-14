from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import StreamingResponse
import json
import xml.etree.ElementTree as ET
from xml.dom import minidom
import io
from docx import Document
from docx.shared import Inches
import logging
from typing import Dict, Any, Union
import re

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="JSON to DOCX Converter API",
    description="Convert JSON files to XML and then to DOCX with character replacements",
    version="1.0.0"
)


class JSONToXMLConverter:
    """Handles conversion from JSON to XML with character replacements."""
    
    def __init__(self):
        self.replacements = {
            '_': ':',
            '$': '@'
        }
    
    def apply_replacements(self, text: str) -> str:
        """Apply character replacements to text."""
        result = text
        for old_char, new_char in self.replacements.items():
            result = result.replace(old_char, new_char)
        return result
    
    def is_xml_content(self, text: str) -> bool:
        """Check if text contains XML content."""
        return isinstance(text, str) and text.strip().startswith('<?xml') and text.strip().endswith('>')
    
    def json_to_xml_element(self, data: Union[Dict, list, str, int, float, bool], 
                           parent_element: ET.Element = None, 
                           element_name: str = "root") -> ET.Element:
        """
        Recursively convert JSON data to XML elements.
        
        Args:
            data: JSON data to convert
            parent_element: Parent XML element
            element_name: Name for the current element
            
        Returns:
            ET.Element: XML element
        """
        # Apply character replacements to element name
        clean_element_name = self.apply_replacements(element_name)
        
        # Ensure element name is valid XML
        clean_element_name = re.sub(r'[^a-zA-Z0-9:@._-]', '_', clean_element_name)
        if clean_element_name and clean_element_name[0].isdigit():
            clean_element_name = f"item_{clean_element_name}"
        
        element = ET.Element(clean_element_name)
        
        if isinstance(data, dict):
            for key, value in data.items():
                child_element = self.json_to_xml_element(value, element, key)
                element.append(child_element)
                
        elif isinstance(data, list):
            for i, item in enumerate(data):
                child_element = self.json_to_xml_element(item, element, f"item_{i}")
                element.append(child_element)
                
        else:
            # For primitive values, apply replacements to the text content
            text_content = str(data)
            element.text = self.apply_replacements(text_content)
        
        return element
    
    def convert_json_to_xml(self, json_data: Union[Dict[str, Any], list]) -> str:
        """
        Convert JSON data to formatted XML string.
        
        Args:
            json_data: Dictionary or list containing JSON data
            
        Returns:
            str: Formatted XML string
        """
        try:
            # Handle special case where JSON contains XML string content
            if isinstance(json_data, list) and len(json_data) == 1:
                item = json_data[0]
                if isinstance(item, dict) and len(item) == 1:
                    key, value = next(iter(item.items()))
                    if self.is_xml_content(value):
                        # Extract and process the XML content directly
                        xml_content = self.apply_replacements(value)
                        return xml_content
            
            # Handle regular JSON to XML conversion
            root_element = self.json_to_xml_element(json_data, element_name="document")
            
            # Create a rough XML string
            rough_string = ET.tostring(root_element, encoding='unicode')
            
            # Pretty print the XML
            reparsed = minidom.parseString(rough_string)
            pretty_xml = reparsed.toprettyxml(indent="  ")
            
            # Remove empty lines
            lines = [line for line in pretty_xml.split('\n') if line.strip()]
            return '\n'.join(lines)
            
        except Exception as e:
            logger.error(f"Error converting JSON to XML: {str(e)}")
            raise HTTPException(status_code=500, detail=f"XML conversion failed: {str(e)}")


class XMLToDocxConverter:
    """Handles conversion from XML to DOCX."""
    
    def create_docx_from_xml(self, xml_content: str) -> io.BytesIO:
        """
        Create a DOCX document from XML content.
        
        Args:
            xml_content: XML content as string
            
        Returns:
            io.BytesIO: DOCX file as bytes
        """
        try:
            # Create a new Document
            doc = Document()
            
            # Add title
            title = doc.add_heading('XML Document Content', 0)
            
            # Try to parse and format XML content
            try:
                # Parse the XML string
                root = ET.fromstring(xml_content)
                
                # Add XML content in a more readable format
                self._add_xml_elements_to_document(doc, root, level=0)
                
            except ET.ParseError as e:
                logger.warning(f"XML parsing failed: {str(e)}, adding raw content")
                # If XML parsing fails, add raw content with proper formatting
                doc.add_paragraph("Raw XML Content:")
                
                # Split into lines and add each line as a paragraph for better readability
                lines = xml_content.split('\n')
                for line in lines:
                    if line.strip():  # Skip empty lines
                        p = doc.add_paragraph(line)
                        p.style = 'Normal'
            
            # Save to BytesIO
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            return doc_buffer
            
        except Exception as e:
            logger.error(f"Error creating DOCX from XML: {str(e)}")
            raise HTTPException(status_code=500, detail=f"DOCX creation failed: {str(e)}")
    
    def _add_xml_elements_to_document(self, doc: Document, element: ET.Element, level: int = 0):
        """
        Add XML elements to document with better formatting for WordprocessingML.
        
        Args:
            doc: Document object
            element: XML element
            level: Indentation level
        """
        indent = "  " * level
        
        # For WordprocessingML, we want to show the structure more clearly
        if element.tag:
            # Add opening tag
            tag_text = f"{indent}<{element.tag}"
            
            # Add attributes if they exist
            if element.attrib:
                for attr_name, attr_value in element.attrib.items():
                    tag_text += f' {attr_name}="{attr_value}"'
            
            if element.text and element.text.strip():
                tag_text += f">{element.text.strip()}</{element.tag}>"
                para = doc.add_paragraph(tag_text)
            elif len(element) == 0:
                tag_text += "/>"
                para = doc.add_paragraph(tag_text)
            else:
                tag_text += ">"
                para = doc.add_paragraph(tag_text)
                
                # Add child elements
                for child in element:
                    self._add_xml_elements_to_document(doc, child, level + 1)
                
                # Add closing tag
                closing_para = doc.add_paragraph(f"{indent}</{element.tag}>")
            
            # Set paragraph style
            para.style = 'Normal'


# Initialize converters
json_to_xml_converter = JSONToXMLConverter()
xml_to_docx_converter = XMLToDocxConverter()


@app.post("/convert-json-to-docx/")
async def convert_json_to_docx(file: UploadFile = File(...)):
    """
    Convert uploaded JSON file to DOCX format.
    
    - Accepts JSON file
    - Converts to XML with character replacements (_ -> :, $ -> @)
    - Converts XML to DOCX binary file
    - Returns DOCX file for download
    """
    try:
        # Validate file type
        if not file.filename.endswith('.json'):
            raise HTTPException(
                status_code=400, 
                detail="Only JSON files are accepted"
            )
        
        # Read and parse JSON file
        content = await file.read()
        try:
            json_data = json.loads(content.decode('utf-8'))
        except json.JSONDecodeError as e:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid JSON format: {str(e)}"
            )
        except UnicodeDecodeError as e:
            raise HTTPException(
                status_code=400,
                detail=f"File encoding error: {str(e)}"
            )
        
        logger.info(f"Processing JSON file: {file.filename}")
        
        # Convert JSON to XML
        xml_content = json_to_xml_converter.convert_json_to_xml(json_data)
        logger.info("JSON to XML conversion completed")
        
        # Convert XML to DOCX
        docx_buffer = xml_to_docx_converter.create_docx_from_xml(xml_content)
        logger.info("XML to DOCX conversion completed")
        
        # Prepare response
        output_filename = file.filename.replace('.json', '.docx')
        
        return StreamingResponse(
            io.BytesIO(docx_buffer.getvalue()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={output_filename}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"An unexpected error occurred: {str(e)}"
        )


@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy", "message": "JSON to DOCX Converter API is running"}


@app.get("/")
async def root():
    """Root endpoint with API information."""
    return {
        "message": "JSON to DOCX Converter API",
        "version": "1.0.0",
        "endpoints": {
            "convert": "/convert-json-to-docx/",
            "health": "/health"
        }
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
